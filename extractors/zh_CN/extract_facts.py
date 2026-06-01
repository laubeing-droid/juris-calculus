#!/usr/bin/env python3
"""
juris-calculus 动态多法域事实提取引擎 v1.0
从 raw_bak 脱敏副本中提取结构化的 .facts.json
"""
import os, re, sys, json, datetime
from pathlib import Path
from collections import defaultdict

RAW_BAK = Path(os.environ.get("LEGALOS_DATA", "./data/raw_bak"))
RAW_OUT = Path(os.environ.get("LEGALOS_OUT", "./data/raw"))

def detect_data_type(text: str) -> dict:
    """检测字段数据类型，用于消除脱敏伪阳性
    
    Returns:
        {"type": "NUMERIC_EXISTS"|"HAS_FORMULA"|"TEXT_ONLY"|"EMPTY",
         "has_formula": bool, "has_amount": bool}
    """
    if not text or not text.strip():
        return {"type": "EMPTY", "has_formula": False, "has_amount": False}
    
    has_amount = bool(re.search(r'\[?Amount_\d+\]?|\d+(?:\.\d+)?\s*[万元亿]', text))
    has_formula = bool(re.search(r'LPR|利率|万分之|%|为基数|计算|倍', text))
    has_numeric = bool(re.search(r'\d+', text))
    
    if has_formula and has_amount:
        return {"type": "HAS_FORMULA", "has_formula": True, "has_amount": True}
    elif has_amount:
        return {"type": "NUMERIC_EXISTS", "has_formula": False, "has_amount": True}
    elif has_numeric:
        return {"type": "NUMERIC_EXISTS", "has_formula": has_formula, "has_amount": has_numeric}
    elif text and len(text) > 5:
        return {"type": "TEXT_ONLY", "has_formula": False, "has_amount": False}
    else:
        return {"type": "EMPTY", "has_formula": False, "has_amount": False}

# ═══════════ 待办二：Parser 3.0 图拓扑因果链分析 ═══════════

CAUSAL_PATTERNS = {
    "criminal_nexus": {
        "keywords": ["刑事", "罪", "逮捕", "羁押", "公安", "检察院", "非法", "黑社会",
                     "判决书.*刑事", "刑事.*判决", "侦查", "提起公诉"],
        "weight": 2.5,
        "description": "刑民交叉因果弧"
    },
    "asset_freezing": {
        "keywords": ["查封", "冻结", "保全", "扣押", "财产保全", "诉前保全"],
        "weight": 1.8,
        "description": "资产冻结因果弧"
    },
    "multi_party": {
        "keywords": ["第三人", "共同被告", "连带", "担保人", "保证人", "追加.*被告"],
        "weight": 1.3,
        "description": "多主体混同因果弧"
    },
    "appeal_chain": {
        "keywords": ["二审", "再审", "上诉", "抗诉", "发回重审", "改判", "撤销.*判决"],
        "weight": 1.5,
        "description": "审级递进因果弧"
    },
    "evidence_dispute": {
        "keywords": ["鉴定", "伪造", "虚假", "司法鉴定", "笔迹鉴定", "印章.*伪造"],
        "weight": 2.0,
        "description": "证据争议因果弧"
    },
    "asset_tracing": {
        "keywords": ["房产", "房票", "不动产", "对账", "资金流向", "转账.*记录", 
                     "银行流水", "兑换", "登记.*名下"],
        "weight": 1.8,
        "description": "资产穿透因果弧"
    },
}

def analyze_graph_topology(text: str, domain: str) -> dict:
    """
    Parser 3.0 图拓扑分析：扫描因果链，构建有向弧，增加图密度
    
    替代 Parser 2.0 的关键词计数方式，
    通过检测因果链自动拉高 Base 节点权重。
    """
    arcs = []
    arc_details = {}
    total_weight = 0.0
    base_nodes = 0
    
    # 扫描每种因果弧
    for arc_type, config in CAUSAL_PATTERNS.items():
        count = 0
        samples = []
        for kw in config["keywords"]:
            matches = re.findall(kw, text)
            if matches:
                count += len(matches)
                if len(samples) < 3:
                    samples.append(kw)
        
        if count > 0:
            arcs.append({
                "type": arc_type,
                "description": config["description"],
                "hit_count": count,
                "weight": config["weight"],
                "sample_keywords": samples
            })
            total_weight += count * config["weight"]
            arc_details[arc_type] = {"count": count, "weight": config["weight"]}
    
    # 计算加权节点数
    # 每个因果弧增加 1-3 个等效节点
    arc_count = len(arcs)
    effective_arc_nodes = sum(a["hit_count"] * 0.5 for a in arcs)  # 每2次命中=1个等效节点
    
    # 基础文本节点（段落/章节）
    paragraphs = [p for p in text.split('\n') if len(p.strip()) > 30]
    base_nodes = min(len(paragraphs) // 3, 50)  # 每3段≈1个基础节点，上限50
    
    # 加权 = 基础节点 + 因果弧等效节点 × 域系数
    domain_factor = 1.5 if domain == 'Criminal_or_Administrative' else 1.0
    weighted_nodes = round(base_nodes + effective_arc_nodes * domain_factor, 1)
    
    return {
        "base_nodes": base_nodes,
        "arc_count": arc_count,
        "effective_arc_nodes": round(effective_arc_nodes, 1),
        "weighted_nodes": weighted_nodes,
        "domain_factor": domain_factor,
        "arcs": arcs,
        "parser_version": "3.0"
    }

# ═══════════ 待办三：审级自动识别 ═══════════

STAGE_PATTERNS = {
    "FIRST_INSTANCE": {
        "keywords": ["一审", "起诉状", "立案", "应诉", "民事起诉状", "刑事起诉"],
        "factor": 1.0,
        "label": "一审"
    },
    "APPEAL_RETRIAL": {
        "keywords": ["二审", "上诉", "再审", "抗诉", "发回重审", "民事上诉状", 
                     "刑事上诉状", "再审申请书", "驳回.*上诉", "维持原判"],
        "factor": 1.25,
        "label": "二审/再审"
    },
    "ARBITRATION": {
        "keywords": ["仲裁", "仲裁庭", "仲裁委", "仲裁员", "仲裁申请", "仲裁裁决"],
        "factor": 1.0,
        "label": "仲裁"
    },
    "ENFORCEMENT": {
        "keywords": ["执行", "强制执行", "执行申请", "查封.*执行", "冻结.*执行",
                     "失信被执行人", "限制消费", "执行异议"],
        "factor": 1.1,
        "label": "执行阶段"
    },
}

def detect_stage(text: str, folder_name: str = "") -> dict:
    """
    审级自动识别：扫描文本中的审级关键词
    
    优先级: 执行 > 再审(二审) > 仲裁 > 一审
    返回 γ_stage 系数供 legalos_pricing.py 联动作矩阵乘法
    """
    combined = text + folder_name
    scores = {}
    
    for stage, config in STAGE_PATTERNS.items():
        score = 0
        hits = []
        for kw in config["keywords"]:
            count = len(re.findall(kw, combined))
            if count > 0:
                score += count
                hits.append(kw)
        scores[stage] = {"score": score, "hits": hits[:5]}
    
    # 按优先级判定
    if scores["ENFORCEMENT"]["score"] >= 2:
        stage = "ENFORCEMENT"
    elif scores["APPEAL_RETRIAL"]["score"] >= 2:
        stage = "APPEAL_RETRIAL"
    elif scores["ARBITRATION"]["score"] >= 2:
        stage = "ARBITRATION"
    else:
        stage = "FIRST_INSTANCE"
    
    config = STAGE_PATTERNS[stage]
    
    return {
        "stage": stage,
        "label": config["label"],
        "gamma_factor": config["factor"],
        "confidence": min(scores[stage]["score"] / 3, 1.0),
        "hit_keywords": scores[stage]["hits"]
    }

# ═══════════ 待办一：DP差分隐私影子节点 ═══════════

def generate_dp_shadow(text: str, domain: str) -> dict:
    """
    差分隐私影子节点生成：
    从文本中提取金额，调用 Laplace 加噪，
    保持本金/利息/罚息比例关系，
    写入 _dp_shadow 节点供公网同步使用。
    """
    # 提取所有金额
    amount_matches = re.findall(r'(\d+(?:,\d{3})*(?:\.\d{1,2})?)\s*(万)?\s*(元|美元|人民币)?', text)
    
    if not amount_matches:
        return None
    
    # 解析为浮点数
    amounts = []
    for num_str, wan, unit in amount_matches:
        try:
            val = float(num_str.replace(',', ''))
            if wan:
                val *= 10000  # 万 → 元
            if val > 100:  # 过滤太小的数字
                amounts.append(val)
        except:
            pass
    
    if len(amounts) < 2:
        return None
    
    # 去重并排序
    amounts = sorted(set(amounts), reverse=True)[:10]
    
    try:
        from legalos_services.differential_privacy import RatioPreservingDP
        dp = RatioPreservingDP(epsilon=1.0)
        
        # 对金额列表加噪（保持比例）
        noisy = dp.noiser.add_noise_batch(amounts, preserve_ratio=True)
        
        # 构建影子节点
        shadow = {
            "epsilon": 1.0,
            "amount_count": len(amounts),
            "original_range": f"{min(amounts):.0f} ~ {max(amounts):.0f}",
            "noisy_range": f"{min(noisy):.0f} ~ {max(noisy):.0f}",
            "ratio_preserved": True,
            "amounts": [{"original": round(a, 2), "noisy": round(n, 2)} 
                       for a, n in zip(amounts[:5], noisy[:5])],
            "usage": "公网同步/云端审计专用。原始金额仅存于本地沙箱。"
        }
        
        # 验证比例保持
        if len(amounts) >= 2 and amounts[1] > 0 and noisy[1] > 0:
            orig_ratio = amounts[0] / amounts[1]
            noisy_ratio = noisy[0] / max(0.01, noisy[1])
            shadow["ratio_error_pct"] = round(abs(orig_ratio - noisy_ratio) / orig_ratio * 100, 4)
        
        return shadow
    except ImportError:
        return {"error": "differential_privacy module not available", "amount_count": len(amounts)}

def extract_text_from_docx(filepath):
    """从docx提取文本"""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
        return '\n'.join(paragraphs)
    except Exception as e:
        return f"[EXTRACT_ERROR: {e}]"

def extract_text_from_pdf(filepath):
    """从pdf提取文本"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:10]:  # 最多10页，防止超大PDF
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return '\n'.join(text_parts)
    except Exception as e:
        return f"[EXTRACT_ERROR: {e}]"

def extract_text(filepath):
    """根据文件类型提取文本"""
    ext = os.path.splitext(str(filepath))[1].lower()
    if ext == '.docx':
        return extract_text_from_docx(str(filepath))
    elif ext == '.pdf':
        return extract_text_from_pdf(str(filepath))
    elif ext in {'.txt', '.md', '.json'}:
        try:
            with open(str(filepath), 'r', encoding='utf-8') as f:
                return f.read()
        except:
            return ""
    return ""

def classify_domain(folder_path):
    """根据文件夹路径分类法域"""
    path_str = str(folder_path).lower()
    if '刑事' in path_str:
        return 'Criminal_or_Administrative', '刑事'
    if '行政' in path_str:
        return 'Criminal_or_Administrative', '行政'
    if '非诉' in path_str:
        # 非诉中的合同类 → Civil_Contract
        return 'Civil_Contract', '非诉'
    if '民事' in path_str:
        return 'Civil_Contract', '民事'
    return 'Civil_Contract', '未知'

def extract_event_date(text, domain, category):
    """提取核心日期"""
    date_patterns = [
        r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
        r'(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?',
    ]
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            try:
                y, mo, d = int(m[0]), int(m[1]), int(m[2])
                if 2020 <= y <= 2026 and 1 <= mo <= 12 and 1 <= d <= 31:
                    dates.append(f"{y}{mo:02d}{d:02d}")
            except:
                pass
    
    if dates:
        # 返回最早日期作为事件日
        return sorted(dates)[0]
    return "UNKNOWN"

def extract_civil_contract_facts(text, folder_name):
    """民事合同类事实提取 — Parser 3.0 图拓扑优先"""
    facts = {
        "domain": "Civil_Contract",
        "event_date": extract_event_date(text, "Civil_Contract", "民事"),
        "party_names": [],
        "core_elements": {
            "lease_term_or_delivery": "",
            "payment_rule": "",
            "deposit": ""
        },
        "rigid_clauses": {
            "liquidated_damages": "",
            "force_majeure": "",
            "dispute_resolution": ""
        }
    }
    
    # ═══ Parser 3.0: 因果箭头优先 ═══
    topo = analyze_graph_topology(text, "Civil_Contract")
    arrow_count = topo["arc_count"]
    causal_weight = topo["weighted_nodes"]
    
    # 当事人名称提取
    facts["party_names"] = ["[甲方]", "[乙方]"]
    if '第三人' in folder_name or '独立第三方' in folder_name:
        facts["party_names"].append("[独立第三方_1]")
    
    # 租赁期限/交付周期
    lease_patterns = [
        r'租赁期限[：:]\s*(.+?)(?:\n|。|；)',
        r'租期[：:]\s*(.+?)(?:\n|。|；)',
        r'租赁期间[：:]\s*(.+?)(?:\n|。|；)',
        r'交付[期限|时间][：:]\s*(.+?)(?:\n|。|；)',
        r'交货[期限|周期][：:]\s*(.+?)(?:\n|。|；)',
    ]
    for pat in lease_patterns:
        m = re.search(pat, text)
        if m:
            facts["core_elements"]["lease_term_or_delivery"] = m.group(1).strip()[:200]
            break
    
    # 付款规则
    payment_patterns = [
        r'(?:付款|支付|价款)[方式|条件|规则][：:]\s*(.+?)(?:\n|。|；)',
        r'(?:约定|合同约定)[^。]*?(?:支付|付款)[^。]*?(?:。|；)',
    ]
    for pat in payment_patterns:
        m = re.search(pat, text)
        if m:
            facts["core_elements"]["payment_rule"] = m.group(0).strip()[:200]
            break
    
    # 押金/担保
    deposit_patterns = [
        r'(?:押金|保证金|定金|履约担保)[：:]\s*(.+?)(?:\n|。|；)',
        r'(?:交付|支付|交纳).*?(?:押金|保证金)',
    ]
    for pat in deposit_patterns:
        m = re.search(pat, text)
        if m:
            facts["core_elements"]["deposit"] = m.group(0).strip()[:200]
            break
    
    # 违约金
    ld_patterns = [
        r'违约[责任金].*?[：:，,]\s*(.+?)(?:\n|。|；)',
        r'(?:任何一方|一方|双方).*?违约.*?(?:支付|赔偿|承担).*?(?:违约金|损失).*?(?:。|；)',
    ]
    for pat in ld_patterns:
        m = re.search(pat, text)
        if m:
            facts["rigid_clauses"]["liquidated_damages"] = m.group(0).strip()[:300]
            break
    
    # 不可抗力
    fm_patterns = [
        r'不可抗力[：:]\s*(.+?)(?:\n|。|；)',
        r'(?:因|由于).*?不可抗力.*?(?:免除|不承担|免责).*?(?:。|；)',
    ]
    for pat in fm_patterns:
        m = re.search(pat, text)
        if m:
            facts["rigid_clauses"]["force_majeure"] = m.group(0).strip()[:300]
            break
    
    # 争议解决
    dr_patterns = [
        r'(?:争议|纠纷).*?(?:解决|处理).*?[：:]\s*(.+?)(?:\n|。|；)',
        r'(?:管辖|向).*?(?:法院|仲裁).*?(?:起诉|申请|解决).*?(?:。|；)',
        r'(?:提交|交由).*?(?:法院|仲裁委).*?(?:管辖|审理|裁决).*?(?:。|；)',
    ]
    for pat in dr_patterns:
        m = re.search(pat, text)
        if m:
            facts["rigid_clauses"]["dispute_resolution"] = m.group(0).strip()[:300]
            break
    
    # ═══ Parser 3.0: 因果箭头分析注入 ═══
    facts["_causal_analysis"] = {
        "arrow_count": arrow_count,
        "causal_weight": causal_weight,
        "method": "Parser 3.0 图拓扑优先 - 因果连线（箭头数）替代词汇计数"
    }
    
    return facts

def extract_criminal_admin_facts(text, folder_name):
    """刑事/行政类事实提取"""
    facts = {
        "domain": "Criminal_or_Administrative",
        "event_date": extract_event_date(text, "Criminal_or_Administrative", "刑事/行政"),
        "party_names": [],
        "core_elements": {
            "charges_or_actions": "",
            "key_evidence_list": ""
        },
        "rigid_clauses": {
            "procedural_timeline": "",
            "statutory_basis": ""
        }
    }
    
    # 当事人
    if '刑事' in str(folder_name):
        facts["party_names"] = ["[涉案人_1]", "[办案单位]"]
    else:
        facts["party_names"] = ["[行政相对人]", "[行政主体]"]
    
    # 涉嫌罪名/行政行为
    charge_patterns = [
        r'涉嫌[：:]?\s*(.+?罪)(?:\n|。|；|，)',
        r'(?:罪名|涉嫌)[：:]\s*(.+?)(?:\n|。|；)',
        r'(?:行政处罚|行政决定)[：:]\s*(.+?)(?:\n|。|；)',
        r'(?:认定|查明|确认)[^。]*?(?:违法|违规).*?(?:。|；)',
    ]
    for pat in charge_patterns:
        m = re.search(pat, text)
        if m:
            facts["core_elements"]["charges_or_actions"] = m.group(0).strip()[:300]
            break
    
    # 证据链
    evidence_keywords = ['证人证言', '鉴定意见', '现场笔录', '书证', '物证', '电子数据',
                        '勘验笔录', '检查笔录', '视听资料', '被告人供述', '被害人陈述']
    evidence_found = []
    for kw in evidence_keywords:
        if kw in text:
            evidence_found.append(kw)
    if evidence_found:
        facts["core_elements"]["key_evidence_list"] = '、'.join(evidence_found)[:300]
    
    # 程序时间线
    proc_patterns = [
        r'(?:拘留|逮捕|取保候审|监视居住|立案|侦查|审查起诉)[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
        r'(\d{4}年\d{1,2}月\d{1,2}日)[^。]*?(?:拘留|逮捕|立案)',
    ]
    timeline = []
    for pat in proc_patterns:
        ms = re.findall(pat, text)
        for m in ms:
            if isinstance(m, tuple):
                timeline.append(''.join(m))
            else:
                timeline.append(m)
    if timeline:
        facts["rigid_clauses"]["procedural_timeline"] = '; '.join(timeline[:5])[:300]
    
    # 法条依据
    law_patterns = [
        r'(?:依照|根据|依据|适用)\s*《(.+?)》\s*第\s*(\d+)\s*条',
        r'《(.+?)》第(\d+)条',
    ]
    laws = []
    for pat in law_patterns:
        ms = re.findall(pat, text)
        for m in ms:
            if isinstance(m, tuple) and len(m) >= 2:
                laws.append(f"《{m[0]}》第{m[1]}条")
    if laws:
        facts["rigid_clauses"]["statutory_basis"] = '; '.join(laws[:5])[:300]
    
    return facts

def process_case_folder(folder_path):
    """处理单个案卷文件夹"""
    folder_name = folder_path.name
    domain, category = classify_domain(str(folder_path))
    
    print(f"\n{'='*60}")
    print(f"[CASE] {folder_name}")
    print(f"  Domain: {domain} | Category: {category}")
    
    # 收集所有文档文本
    all_text = ""
    key_files = []
    
    # 优先找关键文档
    priority_keywords = ['起诉状', '判决', '裁定', '合同', '协议', '答辩', '代理词', 
                        '辩护词', '上诉状', '申请书', '法律意见', '证据清单', '阅卷']
    
    skip_exts = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.mp3', '.mp4', '.m4a', '.wav', '.zip'}
    process_exts = {'.docx', '.pdf', '.txt', '.md', '.json'}
    
    for root, dirs, files in os.walk(str(folder_path)):
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1].lower()
            if ext in skip_exts:
                continue
            if ext not in process_exts:
                continue
            
            filepath = os.path.join(root, fname)
            
            priority = 0
            for kw in priority_keywords:
                if kw in fname:
                    priority = 10
                    break
            
            text = extract_text(filepath)
            if text and len(text) > 50:
                key_files.append((priority, fname, text))
                all_text += text + '\n---PAGE_BREAK---\n'
    
    # 按优先级排序
    key_files.sort(key=lambda x: x[0], reverse=True)
    
    print(f"  Key documents extracted: {len(key_files)}")
    for _, fname, _ in key_files[:5]:
        print(f"    - {fname}")
    
    # 根据法域提取事实
    if domain == 'Civil_Contract':
        facts = extract_civil_contract_facts(all_text, folder_name)
    else:
        facts = extract_criminal_admin_facts(all_text, folder_name)
    
    # ═══ Parser 3.0 升级：图拓扑因果链扫描 ═══
    graph_topo = analyze_graph_topology(all_text, domain)
    facts["_graph_topology"] = graph_topo
    
    # ═══ 待办三：审级自动识别 ═══
    stage_info = detect_stage(all_text, folder_name)
    facts["_stage"] = stage_info
    
    # 生成输出文件名
    case_short = folder_name[:40].replace('/', '_').replace('\\', '_')
    date_str = datetime.datetime.now().strftime('%Y%m%d')
    out_name = f"{date_str}_{case_short}.facts.json"
    out_path = RAW_OUT / out_name
    
    # 附加元数据
    facts["_meta"] = {
        "source_folder": str(folder_path.relative_to(RAW_BAK)),
        "extraction_date": date_str,
        "category": category,
        "stage": stage_info["stage"],
        "documents_processed": len(key_files),
        "total_text_chars": len(all_text)
    }
    
    # 附加数据类型标记
    if domain == 'Civil_Contract':
        facts["_data_types"] = {
            "lease_term_or_delivery": detect_data_type(facts["core_elements"].get("lease_term_or_delivery", "")),
            "payment_rule": detect_data_type(facts["core_elements"].get("payment_rule", "")),
            "deposit": detect_data_type(facts["core_elements"].get("deposit", "")),
            "liquidated_damages": detect_data_type(facts["rigid_clauses"].get("liquidated_damages", "")),
            "force_majeure": detect_data_type(facts["rigid_clauses"].get("force_majeure", "")),
            "dispute_resolution": detect_data_type(facts["rigid_clauses"].get("dispute_resolution", "")),
        }
    else:
        facts["_data_types"] = {
            "charges_or_actions": detect_data_type(facts["core_elements"].get("charges_or_actions", "")),
            "key_evidence_list": detect_data_type(facts["core_elements"].get("key_evidence_list", "")),
            "procedural_timeline": detect_data_type(facts["rigid_clauses"].get("procedural_timeline", "")),
            "statutory_basis": detect_data_type(facts["rigid_clauses"].get("statutory_basis", "")),
        }
    
    # ═══ 待办一：DP差分隐私影子节点 ═══
    dp_shadow = generate_dp_shadow(all_text, domain)
    if dp_shadow:
        facts["_dp_shadow"] = dp_shadow
    
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(facts, f, ensure_ascii=False, indent=2)
    
    print(f"  [OUTPUT] {out_name}")
    print(f"  Facts: domain={facts['domain']}, date={facts['event_date']}, stage={stage_info['stage']}")
    if graph_topo.get('causal_arcs'):
        print(f"  图拓扑: base={graph_topo['base_nodes']}→weighted={graph_topo['weighted_nodes']} ({graph_topo['arc_count']}条因果弧)")
    if dp_shadow:
        print(f"  DP影子: {dp_shadow.get('amount_count',0)}个金额已加噪 (ε={dp_shadow.get('epsilon')})")
    
    return facts, out_path

def main():
    print("=" * 60)
    print("juris-calculus 动态多法域事实提取引擎 v1.0")
    print("=" * 60)
    
    # 确保输出目录存在
    RAW_OUT.mkdir(parents=True, exist_ok=True)
    
    results = []
    
    # 遍历四大分区
    for category_dir in RAW_BAK.iterdir():
        if not category_dir.is_dir():
            continue
        
        print(f"\n>>> Processing category: {category_dir.name}")
        
        for case_folder in category_dir.iterdir():
            if case_folder.is_dir():
                try:
                    facts, out_path = process_case_folder(case_folder)
                    results.append({
                        "folder": str(case_folder.relative_to(RAW_BAK)),
                        "domain": facts["domain"],
                        "output": str(out_path.name)
                    })
                except Exception as e:
                    print(f"  [ERROR] {case_folder.name}: {e}")
    
    # 生成汇总
    print("\n" + "=" * 60)
    print("事实提取完成报告")
    print("=" * 60)
    for r in results:
        print(f"  [{r['domain']}] {r['folder']} → {r['output']}")
    print(f"\n  总计: {len(results)} 个 .facts.json 已生成到 {RAW_OUT}")
    print("=" * 60)
    
    return results

if __name__ == '__main__':
    main()
