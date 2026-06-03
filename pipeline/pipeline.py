#!/usr/bin/env python3
"""
juris-calculus 端到端推理管线 v1.0
输入：案卷文件夹路径 → 输出：推理报告 + 精算报告

用法：
    python pipeline/pipeline.py --case "./data/cases/example"
    python pipeline/pipeline.py --batch "./data/cases/civil"
"""
import sys, os, json, time, re
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, field, asdict

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from compiler_core.types import LegalFact, IRState, LegalDomain
from compiler_core.evaluator import FixpointEvaluator, CriticalClarityFailure, load_rules_from_yaml
from compiler_core.domain_config import get_domain_config

# ── 配置 ──
RULES_YAML = os.path.join(os.path.dirname(__file__), '..', 'configs', 'zh_CN', 'rules.yaml')
REPORT_DIR = os.path.join(os.path.dirname(__file__), '..', 'reports')
os.makedirs(REPORT_DIR, exist_ok=True)

# 加载引擎（全局单例）
ZH_RULES = load_rules_from_yaml(RULES_YAML)
import yaml
_cfg_path = os.path.join(os.path.dirname(__file__), '..', 'configs', 'zh_CN', 'domain_config.yaml')
ZH_CONFIG = get_domain_config(LegalDomain.CIVIL)
if os.path.exists(_cfg_path):
    _cfg = yaml.safe_load(open(_cfg_path, encoding='utf-8'))
    if 'alpha_calibrated' in _cfg:
        ZH_CONFIG.alpha = _cfg['alpha_calibrated']
ENGINE = FixpointEvaluator(ZH_RULES, ZH_CONFIG)


# ── 多租户动态路由器 ──
ONTOLOGY_MAP_PATH = os.path.join(os.path.dirname(__file__), '..', 'configs', 'zh_CN', 'ontology_map.yaml')
_ONTOLOGY = yaml.safe_load(open(ONTOLOGY_MAP_PATH, encoding='utf-8')) if os.path.exists(ONTOLOGY_MAP_PATH) else {}

def load_isolated_ontology(case_text: str) -> tuple:
    """按案卷文本锁定唯一命名空间，返回 (域名, 原子注册表)"""
    for domain_name, domain_cfg in _ONTOLOGY.items():
        if not isinstance(domain_cfg, dict) or 'regex_triggers' not in domain_cfg:
            continue
        if any(trigger in case_text for trigger in domain_cfg['regex_triggers']):
            return domain_name, domain_cfg.get('fact_atoms', {})

    # fallback: 默认合同域
    cc = _ONTOLOGY.get('civil_contract', {})
    return 'civil_contract', cc.get('fact_atoms', {})

def filter_rules_by_namespace(facts: dict, case_text: str = "") -> list:
    """按案卷文本匹配 ontology 的 regex_triggers 锁定规则域"""
    domain, _ = load_isolated_ontology(case_text)
    NS_MAP = {
        'contract': ['contract'],
        'tort': ['tort'],
        'criminal': ['criminal'],
        'administrative': ['admin'],
        'corporate': ['corporate'],
        'intellectual_property': ['ip'],
        'family': ['family'],
        'property': ['tort'],
        'state_compensation': ['admin'],
        'enforcement': ['enforcement'],
        'procedure': ['procedure'],
        'juvenile': ['juvenile'],
        'international': ['contract'],
    }
    allowed = NS_MAP.get(domain, ['general']) + ['general']
    return [r for r in ZH_RULES if getattr(r, 'namespace', 'general') in allowed]


@dataclass
class PipelineResult:
    case_id: str
    status: str = "OK"  # OK / HALTED / ERROR
    claims_found: int = 0
    deterministic: int = 0
    tainted: int = 0
    critical: int = 0
    convergence: bool = False
    taint_depth: int = 0
    elapsed_ms: float = 0.0
    pred_hours: float = 0.0
    top_claims: List[dict] = field(default_factory=list)
    ocr_refs: List[dict] = field(default_factory=list)
    blocked_reasons: List[str] = field(default_factory=list)
    trace: str = ""
    error: str = ""


def extract_text_from_case(case_path: str) -> str:
    """从案卷文件夹提取全文（按扩展名扫描）"""
    case_dir = Path(case_path)
    if not case_dir.exists():
        return ""

    combined = ""
    seen_texts = set()
    max_chars = 50000  # 每案卷最多取5万字
    ocr_engine = None  # 惰性加载

    def _ocr(img_path: str) -> str:
        """惰性加载 OCR + 识别图片文字"""
        nonlocal ocr_engine
        if ocr_engine is None:
            try:
                from paddleocr import PaddleOCR
                ocr_engine = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
            except ImportError:
                return ""
        try:
            result = ocr_engine.ocr(img_path, cls=True)
            texts = []
            for line in result or []:
                for word in line or []:
                    txt = word[1][0] if len(word) > 1 else ""
                    if txt and txt.strip():
                        texts.append(txt.strip())
            return '\n'.join(texts)
        except Exception:
            return ""

    # 按扩展名优先级扫描
    ext_priority = ['.docx', '.txt', '.md', '.json', '.pdf', '.doc', '.jpg', '.jpeg', '.png', '.bmp']

    for ext in ext_priority:
        for fp in sorted(case_dir.rglob(f'*{ext}')):
            if fp.is_dir() or fp.name.startswith('~$'):  # 跳过临时文件
                continue
            try:
                text = ""
                if ext == '.docx':
                    from docx import Document
                    doc = Document(str(fp))
                    text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                elif ext == '.doc':
                    # 尝试用 python-docx 兼容读取
                    try:
                        from docx import Document
                        doc = Document(str(fp))
                        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                    except Exception:
                        text = f"[不支持读取 .doc 旧格式: {fp.name}]"
                elif ext == '.pdf':
                    try:
                        import pdfplumber
                        with pdfplumber.open(str(fp)) as pdf:
                            texts = []
                            for page in pdf.pages[:5]:  # 最多5页
                                t = page.extract_text() or ""
                                texts.append(t)
                            text = '\n'.join(texts)
                        # 如果 pdfplumber 没提取到文字（扫描件），用 OCR
                        if len(text.strip()) < 50:
                            ocr_text = _ocr(str(fp))
                            if ocr_text:
                                text = ocr_text
                    except ImportError:
                        text = f"[pdfplumber 未安装]"
                    except Exception:
                        text = f"[PDF解析失败: {fp.name}]"
                elif ext in ('.txt', '.md', '.json'):
                    text = fp.read_text(encoding='utf-8', errors='ignore')
                elif ext in ('.jpg', '.jpeg', '.png', '.bmp'):
                    text = _ocr(str(fp))

                if text and len(text.strip()) > 50:
                    sig = text[:200]
                    if sig not in seen_texts:
                        seen_texts.add(sig)
                        combined += f"\n--- {fp.name} ---\n{text}\n"
            except Exception:
                pass

    return combined


# ── 法律文本正规化看门狗 ──
LEGAL_NORMALIZE = {
    "主动投案": "自首",
    "如实供述": "自首",
    "自动投案": "自首",
    "要求法庭从轻": "申请调减",
    "请求从轻处罚": "申请调减",
    "喝多了不记得": "酒后失去行为能力",
    "一时冲动": "临时起意",
    "不知情": "不明知",
    "没想那么多": "疏忽大意",
    "经常找我要钱": "敲诈勒索",
    "打了几下": "殴打",
    "拿走了": "窃取",
    "骗了": "诈骗",
    "分几次给": "分期交付",
    "马上还": "还款承诺",
    "保证还": "还款承诺",
    "发信息威胁": "恐吓",
    "跟踪": "尾随跟踪",
    "推了一把": "推搡",
    "摔倒了": "倒地",
    "砍了": "持械攻击",
    "捅了": "持械刺伤",
}

def normalize_legal_text(text: str) -> str:
    """正规化看门狗：口语/笔误 → 标准法律术语"""
    for raw, standard in LEGAL_NORMALIZE.items():
        text = text.replace(raw, standard)
    return text


def fact_predicates_from_text(text: str) -> Dict[str, str]:
    """从文本中提取事实谓词（正则版，供 pipeline 使用）"""
    facts = {}
    patterns = [
        (r'(?:民间借贷|借款|借条|欠条)', 'loan_contract'),
        (r'(?:买卖|购销|采购)\s*(?:合同)', 'sales_contract'),
        (r'(?:租赁|租用|出租|承租)', 'lease_contract'),
        (r'(?:建设工程|施工|工程款|总包)', 'construction_contract'),
        (r'(?:物业服务|物业费|物业合同)', 'property_contract'),
        (r'(?:劳动合同|劳务|劳动关.)', 'labor_contract'),
        (r'(?:保险合同|投保|理赔|保险金)', 'insurance_contract'),
        (r'(?:转让|股权|股份)', 'equity_transfer'),
        (r'(?:担保|保证|抵押|质押|保证人)', 'security_contract'),
        (r'(?:合同|协议)\s*(?:成立|签署|生效|已签)', 'contract_formed'),
        (r'(?:合同|协议)\s*(?:无效|撤销|解除|终止)', 'contract_invalid'),
        (r'(?:违约|违反|不履行|拖欠)', 'breach_alleged'),
        (r'(?:未|没有|拒不|逾期|迟迟未)\s*(?:支付|付款|还款)', 'payment_default'),
        (r'(?:已|已经)?\s*(?:交付|交货|发货|提供|给付)', 'goods_delivered'),
        (r'(?:到期|届满|期限届满|已到期|应支付)', 'payment_due'),
        (r'(?:利息|违约金|滞纳金|罚息)', 'liquidated_damages'),
        (r'(?:超过|已过|经过)\s*(?:诉讼时效|除斥期间)', 'statute_barred'),
        (r'(?:不可抗力|疫情|自然灾害|政府行为)', 'force_majeure'),
        (r'(?:交通事故|车祸|肇事|机动车)', 'traffic_accident'),
        (r'(?:医疗损害|医疗事故|误诊|手术|医院)', 'medical_malpractice'),
        (r'(?:故意|过失|明知|应知|疏忽)', 'fault_element'),
        (r'(?:损失|损害|伤害|死亡|伤残|医疗费|误工费)', 'damages_suffered'),
        (r'(?:精神损害|精神抚慰|精神赔偿)', 'emotional_distress'),
        (r'(?:查封|冻结|扣押|保全|强制执行)', 'enforcement_action'),
        (r'(?:二审|上诉|再审|抗诉|发回重审)', 'appeal_proceeding'),
        (r'(?:鉴定|评估|审计|检测|检验|勘验)', 'expert_evidence'),
        (r'(?:自首|立功|坦白|认罪|悔罪|退赃|退赔)', 'leniency_factor'),
        (r'(?:贪污|受贿|挪用公款|职务犯罪)', 'corruption_crime'),
        (r'(?:诈骗|集资诈骗|贷款诈骗|合同诈骗)', 'fraud_crime'),
        (r'(?:毒品|贩卖|运输|制造|非法持有)', 'drug_crime'),
        (r'(?:行政|处罚|许可|强制|复议|诉讼)', 'administrative_action'),
        (r'(?:国家赔偿|赔偿义务机关)', 'state_compensation'),
        (r'(?:未成年人|儿童|少年)\s*(?:犯罪|侵害)', 'juvenile_case'),
        (r'(?:不动产|房产|房屋|土地|产权|登记)', 'real_estate'),
        (r'(?:破产|重整|和解|清算)', 'bankruptcy'),
        (r'(?:专利|发明|实用新型|外观设计)', 'patent_dispute'),
        (r'(?:商标|品牌|标识|侵权|假冒)', 'trademark_dispute'),
        (r'(?:著作权|版权|作品|软件|抄袭)', 'copyright_dispute'),
        # ── 法律对抗事实（Horn 阻却原子）──
        (r'(?:催收.{0,2}(?:函|通知|短信)|催款.{0,2}(?:函|通知)|催告)', 'limitation_interrupted'),
        (r'(?:承认.{0,2}债务|确认.{0,2}欠款|还款承诺|确认债务)', 'limitation_interrupted'),
        (r'(?:部分.{0,2}(?:清偿|履行|还款)|还款.{0,2}(?:计划|安排|协商))', 'limitation_interrupted'),
        (r'(?:被告.{0,4}(?:抗辩|主张|认为).{0,10}(?:违约|过高|无效))', 'defendant_requests_reduction'),
        (r'(?:申请.{0,2}(?:调减|酌减)|请求.{0,2}(?:调减|酌减)|违约金.{0,4}过高|违约金.{0,4}(?:调|酌|减))', 'defendant_requests_reduction'),
        (r'(?:显失公平|公序良俗|乘人之危)', 'defense_unconscionable'),
    ]

    for pat, pred_id in patterns:
        if re.search(pat, text):
            # 提取匹配片段作为描述
            m = re.search(pat, text)
            ctx = text[max(0, m.start()-10):m.end()+30].strip()[:100]
            if pred_id not in facts:
                facts[pred_id] = ctx

    return facts


# ── OCR 语义检索兜底 ──
_ocr_index = None  # 惰性加载

def _load_ocr_index():
    """加载 OCR 原文索引"""
    global _ocr_index
    if _ocr_index is not None:
        return _ocr_index
    import pickle, numpy as np
    from pathlib import Path
    idx_dir = Path(os.environ.get("JURIS_OCR_INDEX_DIR", "./data/chroma_db_ocr"))
    emb_path = idx_dir / "embeddings.npy"
    chunk_path = idx_dir / "chunks.pkl"
    if not emb_path.exists() or not chunk_path.exists():
        print("  [OCR索引未找到，跳过语义兜底]")
        _ocr_index = None
        return None
    embeddings = np.load(str(emb_path))
    chunks = pickle.loads(chunk_path.read_bytes())
    _ocr_index = {"embeddings": embeddings, "chunks": chunks}
    print(f"  [OCR索引已加载: {len(chunks)}段]")
    return _ocr_index

def search_ocr(query: str, top_k: int = 3) -> list:
    """用案卷文本搜 OCR 原文"""
    idx = _load_ocr_index()
    if idx is None:
        return []
    from sentence_transformers import SentenceTransformer
    import numpy as np
    model = SentenceTransformer("BAAI/bge-large-zh-v1.5")
    q_emb = model.encode([query[:512]], normalize_embeddings=True)
    scores = np.dot(idx["embeddings"], q_emb.T).flatten()
    top_idx = np.argsort(scores)[-top_k:][::-1]
    results = []
    for i in top_idx:
        chunk = idx["chunks"][i]
        results.append({
            "book": chunk["book"],
            "start_page": chunk["start_page"],
            "end_page": chunk["end_page"],
            "text": chunk["text"][:300],
            "score": round(float(scores[i]), 3),
        })
    return results


def process_case(case_path: str) -> PipelineResult:
    """处理单个案卷"""
    case_dir = Path(case_path)
    case_id = case_dir.name if case_dir.exists() else "unknown"
    result = PipelineResult(case_id=case_id)

    try:
        # 1. 提取文本
        t0 = time.time()
        text = extract_text_from_case(case_path)
        if not text.strip():
            result.status = "ERROR"
            result.error = "无有效文本"
            return result

        # 2. 提取事实谓词（先过正规化看门狗）
        text = normalize_legal_text(text)
        facts = fact_predicates_from_text(text)
        if not facts:
            result.status = "ERROR"
            result.error = "未提取到事实谓词"
            return result

        # 3.5. 前置看门狗 — 0.1ms 硬编码扫射（双向统一注册表）
        try:
            from .prc_us_alignment import run_alignment_watchdog
            # 默认中国人审美国合同：source="US", target="CN"
            wd = run_alignment_watchdog(text, source="US", target="CN")
            if wd["is_blocked"]:
                print(f"  [WATCHDOG] 触发绝对阻断: {wd['block_reasons']}")
                result.blocked_reasons = wd["block_reasons"]
            if wd["pre_triggered_atoms"]:
                facts.update(wd["pre_triggered_atoms"])
        except ImportError:
            pass

        # 3.6. 断言守卫：白名单校验 + 非法原子熔断
        try:
            from .guardian import assert_atoms
            facts, rejected = assert_atoms(facts)
            for rej in rejected:
                print(f"  [GUARDIAN] {rej}")
        except ImportError:
            pass

        # 3. 构建 IRState 并推理（Namespace 路由：合同案卷只看合同规则）
        state = IRState(world_id=case_id)
        for fid, desc in facts.items():
            state.facts[fid] = LegalFact(fid, str(desc)[:200])

        # Namespace 过滤规则
        engine = ENGINE
        filtered_rules = filter_rules_by_namespace(facts, text)
        if len(filtered_rules) < len(ZH_RULES):
            engine = FixpointEvaluator(filtered_rules, ZH_CONFIG)
            engine = FixpointEvaluator(filtered_rules, ZH_CONFIG)

        halted = False
        try:
            state = engine.evaluate(state)
        except CriticalClarityFailure:
            halted = True

        elapsed = round((time.time() - t0) * 1000, 1)
        claims = list(state.claims.values()) if state.claims else []
        det = sum(1 for c in claims if not c.requires_human_review)
        tnt = sum(1 for c in claims if c.requires_human_review and c.confidence >= 0.2)
        cri = sum(1 for c in claims if c.confidence < 0.2)
        max_depth = max((len(c.taint_chain) for c in claims), default=0)

        # 4. 精算
        try:
            from legalos_services.legalos_pricing import LegalOSPricingEngine, PricingCase
            engine = LegalOSPricingEngine()
            ne = max(len(state.facts) + len(claims), 1)
            case_p = PricingCase(effective_nodes=ne, location="LOCAL", stage="FIRST_INSTANCE")
            pricing = engine.predict_hours(case_p)
            pred_hours = pricing['total_hours']
        except Exception:
            pred_hours = 0.0

        # 5. 整理
        top = []
        for c in sorted(claims, key=lambda x: -x.confidence)[:10]:
            top.append({
                "id": c.id[:80],
                "confidence": round(c.confidence, 2),
                "human_review": c.requires_human_review,
                "taint": c.taint_summary()[:100],
            })

        # 6. OCR 语义兜底（结论偏少时补充原文参考）
        ocr_refs = []
        if len(claims) < 30:
            try:
                ocr_refs = search_ocr(text)
            except Exception:
                ocr_refs = []

        result.top_claims = top
        result.ocr_refs = ocr_refs

        result.status = "HALTED" if halted else "OK"
        result.claims_found = len(claims)
        result.deterministic = det
        result.tainted = tnt
        result.critical = cri
        result.convergence = not halted
        result.taint_depth = max_depth
        result.elapsed_ms = elapsed
        result.pred_hours = pred_hours
        result.top_claims = top
        result.trace = f"HALTED after {len(claims)} claims" if halted else f"CONVERGED: {len(claims)} claims"

    except Exception as e:
        result.status = "ERROR"
        result.error = str(e)
        import traceback
        result.trace = traceback.format_exc()

    return result


def export_report(result: PipelineResult, output_path: str):
    """导出推理报告（Markdown）"""
    lines = [
        f"# 案卷推理报告 — {result.case_id}",
        "",
        f"生成时间: {time.strftime('%Y-%m-%d %H:%M')}",
        f"状态: **{result.status}**",
        f"引擎: juris-calculus ({len(ZH_RULES)}条规则)",
        "",
        "## 推理结果",
        "",
        f"| 指标 | 数值 |",
        f"|------|-----:|",
        f"| 结论总数 | {result.claims_found} |",
        f"| 确定性结论 | {result.deterministic} |",
        f"| 污点结论 | {result.tainted} |",
        f"| 低置信度 | {result.critical} |",
        f"| 审计链最大深度 | {result.taint_depth} |",
        f"| 收敛 | {'✅' if result.convergence else '❌'} |",
        f"| 处理耗时 | {result.elapsed_ms:.0f}ms |",
        f"| 预测工时 | {result.pred_hours:.1f}h |",
        "",
    ]

    if result.error:
        lines += ["## 错误", "", f"```\n{result.error}\n```", ""]

    if result.top_claims:
        lines += ["## TOP 10 结论", ""]
        for i, c in enumerate(result.top_claims, 1):
            review = "🔴需人工复核" if c["human_review"] else "✅"
            taint = f" | 审计: {c['taint']}" if c["taint"] != "CLEAR" else ""
            lines.append(f"{i}. [{review}] **{c['id'][:60]}** conf={c['confidence']}{taint}")

    if result.ocr_refs:
        lines += ["", "## OCR 原文参考", ""]
        for i, ref in enumerate(result.ocr_refs, 1):
            lines.append(f"{i}. **[ {ref['book']} / 第{ref['start_page']}-{ref['end_page']}页 ]** (相似度: {ref['score']})")
            lines.append(f"   {ref['text'][:200]}")
            lines.append("")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text('\n'.join(lines), encoding='utf-8')
    print(f"  报告: {output_path}")


def run_single(case_path: str):
    """处理单个案卷"""
    case_name = Path(case_path).name
    print(f"\n{'='*60}")
    print(f"处理案卷: {case_name}")
    print(f"{'='*60}")

    result = process_case(case_path)
    report_path = os.path.join(REPORT_DIR, f"{case_name}_推理报告.md")
    export_report(result, report_path)

    print(f"状态: {result.status}")
    print(f"结论: {result.claims_found} | 确定={result.deterministic} 污点={result.tainted} 低置信={result.critical}")
    print(f"耗时: {result.elapsed_ms:.0f}ms | 预测工时: {result.pred_hours:.1f}h")
    return result


def run_batch(batch_dir: str):
    """批量处理目录下的所有案卷"""
    batch_path = Path(batch_dir)
    if not batch_path.exists():
        print(f"目录不存在: {batch_dir}")
        return

    results = []
    for case_folder in sorted(batch_path.iterdir()):
        if not case_folder.is_dir():
            continue
        # 跳过非案卷目录（如"工作安排.doc"）
        if case_folder.suffix:
            continue
        result = run_single(str(case_folder))
        results.append(result)

    # 汇总
    total = len(results)
    ok = sum(1 for r in results if r.status == "OK")
    halted = sum(1 for r in results if r.status == "HALTED")
    errors = sum(1 for r in results if r.status == "ERROR")
    avg_claims = sum(r.claims_found for r in results) / max(total, 1)
    avg_hours = sum(r.pred_hours for r in results) / max(total, 1)

    summary_path = os.path.join(REPORT_DIR, f"batch_{batch_path.name}_汇总.md")
    lines = [
        f"# 批量推理汇总 — {batch_path.name}",
        "",
        f"生成时间: {time.strftime('%Y-%m-%d %H:%M')}",
        f"案卷数: {total}",
        "",
        f"| 指标 | 数值 |",
        f"|------|-----:|",
        f"| 成功 | {ok} |",
        f"| 诚实拒算 | {halted} |",
        f"| 错误 | {errors} |",
        f"| 平均结论数 | {avg_claims:.1f} |",
        f"| 平均预测工时 | {avg_hours:.1f}h |",
        "",
        "| 案卷 | 状态 | 结论 | 确定 | 污点 | 低置信 | 工时 |",
        "|------|:----:|:----:|:----:|:----:|:-----:|:----:|",
    ]
    for r in results:
        lines.append(f"| {r.case_id} | {r.status} | {r.claims_found} | {r.deterministic} | {r.tainted} | {r.critical} | {r.pred_hours:.1f}h |")

    Path(summary_path).write_text('\n'.join(lines), encoding='utf-8')
    print(f"\n{'='*60}")
    print(f"批量汇总: {summary_path}")
    print(f"  总计: {total} | 成功: {ok} | 拒算: {halted} | 错误: {errors}")
    print(f"  平均: {avg_claims:.1f}结论/案 | {avg_hours:.1f}h/案")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="juris-calculus 端到端推理管线")
    parser.add_argument("--case", help="单案卷路径")
    parser.add_argument("--batch", help="批量案卷目录")
    args = parser.parse_args()

    if args.case:
        run_single(args.case)
    elif args.batch:
        run_batch(args.batch)
    else:
        parser.print_help()
